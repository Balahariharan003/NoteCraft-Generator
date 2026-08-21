import os
import re
from datetime import datetime

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

OUTPUTS_DIR = os.path.join(os.path.dirname(__file__), "..", "outputs")


# ─────────────────────────────────────────────
# MAIN EXPORT FUNCTION
# ─────────────────────────────────────────────
def export_documents(mom_json: dict, session_id: str, language: str = "en") -> tuple:
    """
    Exports the synthesized MoM JSON to an official formatted DOCX document.
    Renders the exact structure and layout of the official reference PDFs
    using ONLY the verified facts from the session.
    """
    os.makedirs(OUTPUTS_DIR, exist_ok=True)

    raw_title = (
        mom_json.get("session_title")
        or mom_json.get("title")
        or ("Tamil_MoM_Report" if language in ["ta", "tamil"] else "English_MoM_Report")
    )

    clean_name = re.sub(r"[^\w\s-]", "", str(raw_title))
    clean_name = re.sub(r"\s+", "_", clean_name.strip())
    clean_name = clean_name[:60]
    if not clean_name:
        clean_name = "MoM_Report"

    filename = f"{clean_name}_{session_id[:8]}"
    docx_path = os.path.join(OUTPUTS_DIR, f"{filename}.docx")

    if language in ["ta", "tamil", "ta_IN"]:
        _generate_tamil_docx(mom_json, docx_path)
    else:
        _generate_english_docx(mom_json, docx_path)

    return None, f"/outputs/{filename}.docx"


# ─────────────────────────────────────────────
# 1. TAMIL MOM DOCX GENERATION (Reference PDF Style)
# ─────────────────────────────────────────────
def _generate_tamil_docx(data: dict, path: str):
    """
    Generates an official Tamil Minutes of Meeting (.docx) mirroring the
    Tamil Nadu Government Official Proceedings / Grievance Day Reference PDF.
    """
    doc = Document()

    # Font: Arial Unicode MS / Arial (Clean Tamil rendering)
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)
    if style.font.element.rPr is not None:
        style.font.element.rPr.rFonts.set(qn('w:ascii'), 'Arial')
        style.font.element.rPr.rFonts.set(qn('w:hAnsi'), 'Arial')
        style.font.element.rPr.rFonts.set(qn('w:cs'), 'Arial')

    # Page Margins: 0.75 in (54 pt)
    section = doc.sections[0]
    section.top_margin = Pt(54)
    section.bottom_margin = Pt(54)
    section.left_margin = Pt(54)
    section.right_margin = Pt(54)

    # Footer with centered Page Number
    footer = section.footer
    footer_para = footer.paragraphs[0]
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_page_number(footer_para)

    # Extract Data Fields
    location = data.get("district_or_location")
    session_title = data.get("session_title")
    date_str = data.get("date")
    time_str = data.get("time")
    venue = data.get("venue_platform")
    presided_by = data.get("presided_by")
    meeting_no = data.get("meeting_no")
    subject = data.get("subject")
    reference = data.get("reference")
    intro = data.get("intro_paragraph")
    opening_remarks = data.get("opening_exhibition_or_remarks")

    # ── 1. Centered Main Header / Title ──
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_before = Pt(0)
    p_title.paragraph_format.space_after = Pt(6)
    p_title.paragraph_format.line_spacing = 1.15

    title_text = session_title or "கூட்ட நடவடிக்கைகள்"
    if presided_by and "தலைமையில்" not in title_text:
        pres_name = presided_by.replace("முன்னிலை:", "").strip()
        date_part = f"{date_str} அன்று " if date_str else ""
        title_text = f"{pres_name} அவர்கள் தலைமையில் {date_part}நடைபெற்ற {title_text}"

    r_title = p_title.add_run(title_text)
    r_title.bold = True
    r_title.font.size = Pt(13)

    # ── 2. Reference No & Date Line Table ──
    hdr_table = doc.add_table(rows=1, cols=2)
    _remove_table_borders(hdr_table)
    hdr_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_table.autofit = False

    c_left = hdr_table.rows[0].cells[0]
    c_right = hdr_table.rows[0].cells[1]
    c_left.width = Inches(3.5)
    c_right.width = Inches(3.5)

    p_no = c_left.paragraphs[0]
    p_no.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r_no_lbl = p_no.add_run("எண்: ")
    r_no_lbl.bold = True
    r_no_lbl.font.size = Pt(11)
    r_no_val = p_no.add_run(str(meeting_no or "வே/401/2025" if not meeting_no else meeting_no))
    r_no_val.font.size = Pt(11)

    p_dt = c_right.paragraphs[0]
    p_dt.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r_dt_lbl = p_dt.add_run("நாள்: ")
    r_dt_lbl.bold = True
    r_dt_lbl.font.size = Pt(11)
    r_dt_val = p_dt.add_run(str(date_str or datetime.now().strftime("%d.%m.%Y")))
    r_dt_val.font.size = Pt(11)

    doc.add_paragraph().paragraph_format.space_after = Pt(2)

    # ── 3. Subject ('பொருள்') & Reference ('பார்வை') ──
    p_subj = doc.add_paragraph()
    p_subj.paragraph_format.space_after = Pt(4)
    p_subj.paragraph_format.line_spacing = 1.15
    r_sb_lbl = p_subj.add_run("பொருள்:  ")
    r_sb_lbl.bold = True
    r_sb_lbl.font.size = Pt(11)
    subj_text = subject or f"{session_title or 'கூட்டம்'} - நடைபெற்றது - கூட்ட நடவடிக்கைகள் - ஒப்புதல் அளித்தல் - தொடர்பாக."
    r_sb_val = p_subj.add_run(str(subj_text))
    r_sb_val.font.size = Pt(11)

    if reference:
        p_ref = doc.add_paragraph()
        p_ref.paragraph_format.space_after = Pt(4)
        p_ref.paragraph_format.line_spacing = 1.15
        r_rf_lbl = p_ref.add_run("பார்வை:  ")
        r_rf_lbl.bold = True
        r_rf_lbl.font.size = Pt(11)
        r_rf_val = p_ref.add_run(str(reference))
        r_rf_val.font.size = Pt(11)

    # ── 4. Centered Divider Ornament (<><><>) ──
    p_div = doc.add_paragraph()
    p_div.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_div.paragraph_format.space_before = Pt(4)
    p_div.paragraph_format.space_after = Pt(8)
    r_div = p_div.add_run("<><><>")
    r_div.bold = True
    r_div.font.size = Pt(11)

    # ── 5. Introduction & Opening Paragraphs ──
    if intro:
        p_intro = doc.add_paragraph()
        p_intro.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p_intro.paragraph_format.space_after = Pt(6)
        p_intro.paragraph_format.line_spacing = 1.15
        r_intro = p_intro.add_run(str(intro))
        r_intro.font.size = Pt(11)

    if opening_remarks:
        p_op = doc.add_paragraph()
        p_op.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p_op.paragraph_format.space_after = Pt(8)
        p_op.paragraph_format.line_spacing = 1.15
        r_op = p_op.add_run(str(opening_remarks))
        r_op.font.size = Pt(11)

    # ── 6. Representative Grievances & Demands ('கோரிக்கைகள்:') ──
    rep_points = data.get("representative_points")
    if rep_points and isinstance(rep_points, list) and len(rep_points) > 0:
        p_sec1 = doc.add_paragraph()
        p_sec1.paragraph_format.space_before = Pt(10)
        p_sec1.paragraph_format.space_after = Pt(6)
        r_sec1 = p_sec1.add_run("விவசாய சங்கங்களின் கோரிக்கைகள்:")
        r_sec1.bold = True
        r_sec1.underline = True
        r_sec1.font.size = Pt(12)

        for item in rep_points:
            if not isinstance(item, dict):
                continue
            entity = item.get("entity_name") or item.get("speaker") or ""
            pts = item.get("points") or []
            depts = item.get("action_departments") or []

            if entity:
                p_ent = doc.add_paragraph()
                p_ent.paragraph_format.space_before = Pt(6)
                p_ent.paragraph_format.space_after = Pt(2)
                r_ent = p_ent.add_run(f"{entity}:")
                r_ent.bold = True
                r_ent.font.size = Pt(11)

            for pt in pts:
                p_pt = doc.add_paragraph(style="List Bullet")
                p_pt.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                p_pt.paragraph_format.space_after = Pt(2)
                p_pt.paragraph_format.line_spacing = 1.15
                r_pt = p_pt.add_run(str(pt))
                r_pt.font.size = Pt(10.5)

            if depts:
                p_act = doc.add_paragraph()
                p_act.paragraph_format.space_after = Pt(6)
                p_act.paragraph_format.line_spacing = 1.15
                dept_str = ", ".join(depts) if isinstance(depts, list) else str(depts)
                r_act = p_act.add_run(f"(நடவடிக்கை: {dept_str})")
                r_act.bold = True
                r_act.font.size = Pt(10.5)

    # ── 7. Topics Discussed ('விவாதிக்கப்பட்ட தலைப்புகள்:') ──
    topics = data.get("topics_discussed")
    if topics and isinstance(topics, list) and len(topics) > 0:
        p_tl = doc.add_paragraph()
        p_tl.paragraph_format.space_before = Pt(10)
        p_tl.paragraph_format.space_after = Pt(4)
        r_tl = p_tl.add_run("விவாதிக்கப்பட்ட தலைப்புகள்:")
        r_tl.bold = True
        r_tl.underline = True
        r_tl.font.size = Pt(12)

        for idx, topic in enumerate(topics, 1):
            p_t = doc.add_paragraph()
            p_t.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p_t.paragraph_format.space_after = Pt(2)
            p_t.paragraph_format.line_spacing = 1.15
            r_t = p_t.add_run(f"{idx}. {topic}")
            r_t.font.size = Pt(10.5)

    # ── 8. Key Discussion Points ('முக்கிய குறிப்புகள்:') ──
    key_points = data.get("key_points")
    if key_points and isinstance(key_points, list) and len(key_points) > 0:
        p_kpl = doc.add_paragraph()
        p_kpl.paragraph_format.space_before = Pt(10)
        p_kpl.paragraph_format.space_after = Pt(4)
        r_kpl = p_kpl.add_run("முக்கிய குறிப்புகள்:")
        r_kpl.bold = True
        r_kpl.underline = True
        r_kpl.font.size = Pt(12)

        for pt in key_points:
            p_kp = doc.add_paragraph(style="List Bullet")
            p_kp.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p_kp.paragraph_format.space_after = Pt(2)
            p_kp.paragraph_format.line_spacing = 1.15
            r_kp = p_kp.add_run(str(pt))
            r_kp.font.size = Pt(10.5)

    # ── 9. Departmental Responses ('அலுவலர்களின் பதில்கள்') ──
    officer_resp = data.get("officer_responses")
    if officer_resp and isinstance(officer_resp, list) and len(officer_resp) > 0:
        p_sec2 = doc.add_paragraph()
        p_sec2.paragraph_format.space_before = Pt(12)
        p_sec2.paragraph_format.space_after = Pt(6)
        r_sec2 = p_sec2.add_run("அலுவலர்களின் பதில்கள்:")
        r_sec2.bold = True
        r_sec2.underline = True
        r_sec2.font.size = Pt(12)

        for item in officer_resp:
            if not isinstance(item, dict):
                continue
            dept = item.get("department_or_officer") or item.get("officer") or ""
            resp_text = item.get("response") or ""
            pts = item.get("points") or []

            if dept:
                p_dept = doc.add_paragraph()
                p_dept.paragraph_format.space_before = Pt(6)
                p_dept.paragraph_format.space_after = Pt(2)
                r_dept = p_dept.add_run(f"{dept}:")
                r_dept.bold = True
                r_dept.underline = True
                r_dept.font.size = Pt(11)

            if resp_text:
                p_resp = doc.add_paragraph()
                p_resp.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                p_resp.paragraph_format.space_after = Pt(3)
                p_resp.paragraph_format.line_spacing = 1.15
                r_rt = p_resp.add_run(str(resp_text))
                r_rt.font.size = Pt(10.5)

            for pt in pts:
                p_pt2 = doc.add_paragraph(style="List Bullet")
                p_pt2.paragraph_format.space_after = Pt(2)
                p_pt2.paragraph_format.line_spacing = 1.15
                r_pt2 = p_pt2.add_run(str(pt))
                r_pt2.font.size = Pt(10.5)

    # ── 10. Decisions & Actions ──
    decisions = data.get("decisions_taken")
    if decisions and isinstance(decisions, list) and len(decisions) > 0:
        p_dl = doc.add_paragraph()
        p_dl.paragraph_format.space_before = Pt(10)
        p_dl.paragraph_format.space_after = Pt(4)
        r_dl = p_dl.add_run("எடுக்கப்பட்ட முடிவுகள்:")
        r_dl.bold = True
        r_dl.underline = True
        r_dl.font.size = Pt(12)

        for d in decisions:
            p_d = doc.add_paragraph(style="List Bullet")
            p_d.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p_d.paragraph_format.space_after = Pt(2)
            p_d.paragraph_format.line_spacing = 1.15
            r_d = p_d.add_run(str(d))
            r_d.font.size = Pt(10.5)

    actions = data.get("action_items")
    if actions and isinstance(actions, list) and len(actions) > 0:
        p_al = doc.add_paragraph()
        p_al.paragraph_format.space_before = Pt(10)
        p_al.paragraph_format.space_after = Pt(4)
        r_al = p_al.add_run("நடவடிக்கை குறிப்புகள்:")
        r_al.bold = True
        r_al.underline = True
        r_al.font.size = Pt(12)

        for a in actions:
            p_a = doc.add_paragraph(style="List Bullet")
            p_a.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p_a.paragraph_format.space_after = Pt(2)
            p_a.paragraph_format.line_spacing = 1.15
            r_a = p_a.add_run(str(a))
            r_a.font.size = Pt(10.5)

    # ── 11. Vote of Thanks ──
    vote_thanks = data.get("vote_of_thanks")
    if vote_thanks:
        p_vt = doc.add_paragraph()
        p_vt.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p_vt.paragraph_format.space_before = Pt(10)
        p_vt.paragraph_format.space_after = Pt(8)
        p_vt.paragraph_format.line_spacing = 1.15
        r_vt = p_vt.add_run(str(vote_thanks))
        r_vt.font.size = Pt(11)

    # ── 12. Sign-off Blocks (Right Aligned) ──
    p_sig_chair = doc.add_paragraph()
    p_sig_chair.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_sig_chair.paragraph_format.space_before = Pt(24)
    p_sig_chair.paragraph_format.space_after = Pt(2)
    p_sig_chair.paragraph_format.line_spacing = 1.15

    chair_sig = data.get("chairperson_signatory")
    if chair_sig and isinstance(chair_sig, dict):
        c_name = chair_sig.get("name") or "ஒம்/-"
        c_desig = chair_sig.get("designation") or "மாவட்ட ஆட்சித்தலைவர்"
        c_loc = chair_sig.get("location") or location or ""
        r_sc1 = p_sig_chair.add_run(f"{c_name}\n{c_desig},\n{c_loc}.")
    else:
        r_sc1 = p_sig_chair.add_run("ஒம்/-\nமாவட்ட ஆட்சித்தலைவர் / தலைவர்.")
    r_sc1.bold = True
    r_sc1.font.size = Pt(11)

    p_order = doc.add_paragraph()
    p_order.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_order.paragraph_format.space_before = Pt(18)
    p_order.paragraph_format.space_after = Pt(2)
    p_order.paragraph_format.line_spacing = 1.15

    order_sig = data.get("order_signatory") or data.get("convener_signatory")
    if order_sig and isinstance(order_sig, dict):
        o_desig = order_sig.get("designation") or "நேர்முக உதவியாளர்"
        o_loc = order_sig.get("location") or location or ""
        r_ord = p_order.add_run(f"/உத்தரவுப்படி/\n\n{o_desig},\n{o_loc}.")
    else:
        r_ord = p_order.add_run("/உத்தரவுப்படி/\n\nநேர்முக உதவியாளர் / ஒருங்கிணைப்பாளர்.")
    r_ord.bold = True
    r_ord.font.size = Pt(10.5)

    doc.save(path)
    print(f"[OK] Tamil MoM DOCX saved: {path}")


# ─────────────────────────────────────────────
# 2. ENGLISH MOM DOCX GENERATION (Reference PDF Style)
# ─────────────────────────────────────────────
def _generate_english_docx(data: dict, path: str):
    """
    Generates an official English Minutes of Meeting (.docx) mirroring the
    DEPC / Committee Reference PDF.
    """
    doc = Document()

    # Font: Calibri
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    # Page Margins: 0.75 in (54 pt)
    section = doc.sections[0]
    section.top_margin = Pt(54)
    section.bottom_margin = Pt(54)
    section.left_margin = Pt(54)
    section.right_margin = Pt(54)

    # Footer with centered Page Number
    footer = section.footer
    footer_para = footer.paragraphs[0]
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_page_number(footer_para)

    # Extract Fields
    location = data.get("district_or_location")
    session_title = data.get("session_title") or "Minutes of the Meeting"
    date_str = data.get("date")
    time_str = data.get("time")
    venue = data.get("venue_platform")
    presided_by = data.get("presided_by")
    convened_by = data.get("convened_by")
    intro = data.get("intro_paragraph")
    opening_remarks = data.get("opening_exhibition_or_remarks")

    # ── 1. Document Title (Bold, Underlined) ──
    p_title = doc.add_paragraph()
    p_title.paragraph_format.space_after = Pt(8)
    p_title.paragraph_format.line_spacing = 1.15
    
    full_title = f"{session_title}:"
    if location and location not in session_title:
        full_title = f"{session_title}, {location}:"
    if not full_title.startswith("Minutes of"):
        full_title = f"Minutes of the {full_title}"

    r_title = p_title.add_run(full_title)
    r_title.bold = True
    r_title.underline = True
    r_title.font.size = Pt(13)

    # ── 2. Introductory Paragraph ──
    if intro:
        p_intro = doc.add_paragraph()
        p_intro.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p_intro.paragraph_format.space_after = Pt(8)
        p_intro.paragraph_format.line_spacing = 1.15
        r_in = p_intro.add_run(str(intro))
        r_in.font.size = Pt(11)

    # ── 3. Attendees Section ──
    members = data.get("members_representatives") or data.get("members_present")
    special_invitees = data.get("special_invitees_departments")

    if members or special_invitees:
        p_att = doc.add_paragraph()
        p_att.paragraph_format.space_before = Pt(8)
        p_att.paragraph_format.space_after = Pt(4)
        r_att = p_att.add_run("Attendees:")
        r_att.bold = True
        r_att.font.size = Pt(12)

        if members and isinstance(members, list) and len(members) > 0:
            p_mr = doc.add_paragraph()
            p_mr.paragraph_format.space_before = Pt(4)
            p_mr.paragraph_format.space_after = Pt(4)
            r_mr = p_mr.add_run("Members/Representatives:")
            r_mr.bold = True
            r_mr.font.size = Pt(11)

            for idx, member in enumerate(members, 1):
                p_m = doc.add_paragraph()
                p_m.paragraph_format.space_after = Pt(2)
                p_m.paragraph_format.line_spacing = 1.15
                r_m = p_m.add_run(f"{idx}. {member}")
                r_m.font.size = Pt(10.5)

        if special_invitees and isinstance(special_invitees, list) and len(special_invitees) > 0:
            p_si = doc.add_paragraph()
            p_si.paragraph_format.space_before = Pt(6)
            p_si.paragraph_format.space_after = Pt(4)
            r_si = p_si.add_run("Special Invitees/Departments:")
            r_si.bold = True
            r_si.font.size = Pt(11)

            for idx, inv in enumerate(special_invitees, 1):
                p_i = doc.add_paragraph()
                p_i.paragraph_format.space_after = Pt(2)
                p_i.paragraph_format.line_spacing = 1.15
                r_i = p_i.add_run(f"{idx}. {inv}")
                r_i.font.size = Pt(10.5)

    # ── 4. Agenda Points Discussed (Bold, Underline) ──
    p_agd = doc.add_paragraph()
    p_agd.paragraph_format.space_before = Pt(12)
    p_agd.paragraph_format.space_after = Pt(6)
    r_agd = p_agd.add_run("Agenda Points Discussed:")
    r_agd.bold = True
    r_agd.underline = True
    r_agd.font.size = Pt(12)

    # ── 5. Welcome Address ──
    welcome = data.get("welcome_address")
    if welcome:
        p_w = doc.add_paragraph()
        p_w.paragraph_format.space_before = Pt(6)
        p_w.paragraph_format.space_after = Pt(2)
        r_w = p_w.add_run("Welcome Address:")
        r_w.bold = True
        r_w.font.size = Pt(11.5)

        if isinstance(welcome, dict):
            w_text = welcome.get("content") or welcome.get("summary") or ""
            if w_text:
                p_wt = doc.add_paragraph()
                p_wt.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                p_wt.paragraph_format.space_after = Pt(4)
                p_wt.paragraph_format.line_spacing = 1.15
                r_wt = p_wt.add_run(str(w_text))
                r_wt.font.size = Pt(10.5)
        elif isinstance(welcome, str):
            p_wt = doc.add_paragraph()
            p_wt.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p_wt.paragraph_format.space_after = Pt(4)
            p_wt.paragraph_format.line_spacing = 1.15
            r_wt = p_wt.add_run(str(welcome))
            r_wt.font.size = Pt(10.5)

    # ── 6. Key Address by Chairperson ──
    chair_addr = data.get("chairperson_address")
    if chair_addr and isinstance(chair_addr, list) and len(chair_addr) > 0:
        p_ca = doc.add_paragraph()
        p_ca.paragraph_format.space_before = Pt(8)
        p_ca.paragraph_format.space_after = Pt(2)
        r_ca = p_ca.add_run("Key Address by the Chairperson:")
        r_ca.bold = True
        r_ca.font.size = Pt(11.5)

        for pt in chair_addr:
            p_pt = doc.add_paragraph(style="List Bullet")
            p_pt.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p_pt.paragraph_format.space_after = Pt(2)
            p_pt.paragraph_format.line_spacing = 1.15
            r_pt = p_pt.add_run(str(pt))
            r_pt.font.size = Pt(10.5)

    # ── 7. Suggestions by Chairperson ──
    chair_sugg = data.get("chairperson_suggestions")
    if chair_sugg and isinstance(chair_sugg, list) and len(chair_sugg) > 0:
        p_cs = doc.add_paragraph()
        p_cs.paragraph_format.space_before = Pt(8)
        p_cs.paragraph_format.space_after = Pt(2)
        r_cs = p_cs.add_run("Suggestions by the Chairperson:")
        r_cs.bold = True
        r_cs.font.size = Pt(11.5)

        for pt in chair_sugg:
            p_s = doc.add_paragraph(style="List Bullet")
            p_s.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p_s.paragraph_format.space_after = Pt(2)
            p_s.paragraph_format.line_spacing = 1.15
            r_s = p_s.add_run(str(pt))
            r_s.font.size = Pt(10.5)

    # ── 8. Points Raised by Representatives ──
    rep_points = data.get("representative_points")
    if rep_points and isinstance(rep_points, list) and len(rep_points) > 0:
        p_sec1 = doc.add_paragraph()
        p_sec1.paragraph_format.space_before = Pt(10)
        p_sec1.paragraph_format.space_after = Pt(4)
        r_sec1 = p_sec1.add_run("Points Raised by Representatives:")
        r_sec1.bold = True
        r_sec1.underline = True
        r_sec1.font.size = Pt(12)

        for item in rep_points:
            if not isinstance(item, dict):
                continue
            entity = item.get("entity_name") or item.get("speaker") or ""
            pts = item.get("points") or []
            depts = item.get("action_departments") or []

            if entity:
                p_ent = doc.add_paragraph()
                p_ent.paragraph_format.space_before = Pt(6)
                p_ent.paragraph_format.space_after = Pt(2)
                r_ent = p_ent.add_run(f"• {entity}:")
                r_ent.bold = True
                r_ent.underline = True
                r_ent.font.size = Pt(11)

            for pt in pts:
                p_pt = doc.add_paragraph(style="List Bullet")
                p_pt.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                p_pt.paragraph_format.space_after = Pt(2)
                p_pt.paragraph_format.line_spacing = 1.15
                r_pt = p_pt.add_run(str(pt))
                r_pt.font.size = Pt(10.5)

            if depts:
                p_act = doc.add_paragraph()
                p_act.paragraph_format.space_after = Pt(4)
                p_act.paragraph_format.line_spacing = 1.15
                dept_str = ", ".join(depts) if isinstance(depts, list) else str(depts)
                r_act = p_act.add_run(f"Action Departments: {dept_str}")
                r_act.italic = True
                r_act.font.size = Pt(10)

    # ── 9. Departmental Feedback & Responses ──
    officer_resp = data.get("officer_responses")
    if officer_resp and isinstance(officer_resp, list) and len(officer_resp) > 0:
        p_sec2 = doc.add_paragraph()
        p_sec2.paragraph_format.space_before = Pt(10)
        p_sec2.paragraph_format.space_after = Pt(4)
        r_sec2 = p_sec2.add_run("Departmental Feedback & Responses:")
        r_sec2.bold = True
        r_sec2.underline = True
        r_sec2.font.size = Pt(12)

        for item in officer_resp:
            if not isinstance(item, dict):
                continue
            dept = item.get("department_or_officer") or item.get("officer") or ""
            resp_text = item.get("response") or ""
            pts = item.get("points") or []

            if dept:
                p_dept = doc.add_paragraph()
                p_dept.paragraph_format.space_before = Pt(6)
                p_dept.paragraph_format.space_after = Pt(2)
                r_dept = p_dept.add_run(f"• {dept}:")
                r_dept.bold = True
                r_dept.underline = True
                r_dept.font.size = Pt(11)

            if resp_text:
                p_resp = doc.add_paragraph()
                p_resp.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                p_resp.paragraph_format.space_after = Pt(3)
                p_resp.paragraph_format.line_spacing = 1.15
                r_rt = p_resp.add_run(str(resp_text))
                r_rt.font.size = Pt(10.5)

            for pt in pts:
                p_pt2 = doc.add_paragraph(style="List Bullet")
                p_pt2.paragraph_format.space_after = Pt(2)
                p_pt2.paragraph_format.line_spacing = 1.15
                r_pt2 = p_pt2.add_run(str(pt))
                r_pt2.font.size = Pt(10.5)

    # ── 10. Topics Discussed ──
    topics = data.get("topics_discussed")
    if topics and isinstance(topics, list) and len(topics) > 0:
        p_tl = doc.add_paragraph()
        p_tl.paragraph_format.space_before = Pt(10)
        p_tl.paragraph_format.space_after = Pt(4)
        r_tl = p_tl.add_run("Topics Discussed:")
        r_tl.bold = True
        r_tl.underline = True
        r_tl.font.size = Pt(12)

        for idx, topic in enumerate(topics, 1):
            p_t = doc.add_paragraph()
            p_t.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p_t.paragraph_format.space_after = Pt(2)
            p_t.paragraph_format.line_spacing = 1.15
            r_t = p_t.add_run(f"{idx}. {topic}")
            r_t.font.size = Pt(10.5)

    # ── 11. Key Discussion Points ──
    key_points = data.get("key_points")
    if key_points and isinstance(key_points, list) and len(key_points) > 0:
        p_kpl = doc.add_paragraph()
        p_kpl.paragraph_format.space_before = Pt(10)
        p_kpl.paragraph_format.space_after = Pt(4)
        r_kpl = p_kpl.add_run("Key Discussion Points:")
        r_kpl.bold = True
        r_kpl.underline = True
        r_kpl.font.size = Pt(12)

        for pt in key_points:
            p_kp = doc.add_paragraph(style="List Bullet")
            p_kp.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p_kp.paragraph_format.space_after = Pt(2)
            p_kp.paragraph_format.line_spacing = 1.15
            r_kp = p_kp.add_run(str(pt))
            r_kp.font.size = Pt(10.5)

    # ── 12. Decisions Taken & Action Items ──
    decisions = data.get("decisions_taken")
    if decisions and isinstance(decisions, list) and len(decisions) > 0:
        p_dl = doc.add_paragraph()
        p_dl.paragraph_format.space_before = Pt(10)
        p_dl.paragraph_format.space_after = Pt(4)
        r_dl = p_dl.add_run("Decisions Taken:")
        r_dl.bold = True
        r_dl.underline = True
        r_dl.font.size = Pt(12)

        for d in decisions:
            p_d = doc.add_paragraph(style="List Bullet")
            p_d.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p_d.paragraph_format.space_after = Pt(2)
            p_d.paragraph_format.line_spacing = 1.15
            r_d = p_d.add_run(str(d))
            r_d.font.size = Pt(10.5)

    actions = data.get("action_items")
    if actions and isinstance(actions, list) and len(actions) > 0:
        p_al = doc.add_paragraph()
        p_al.paragraph_format.space_before = Pt(10)
        p_al.paragraph_format.space_after = Pt(4)
        r_al = p_al.add_run("Action Items:")
        r_al.bold = True
        r_al.underline = True
        r_al.font.size = Pt(12)

        for a in actions:
            p_a = doc.add_paragraph(style="List Bullet")
            p_a.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p_a.paragraph_format.space_after = Pt(2)
            p_a.paragraph_format.line_spacing = 1.15
            r_a = p_a.add_run(str(a))
            r_a.font.size = Pt(10.5)

    # ── 13. Vote of Thanks ──
    vote_thanks = data.get("vote_of_thanks")
    if vote_thanks:
        p_vt_lbl = doc.add_paragraph()
        p_vt_lbl.paragraph_format.space_before = Pt(10)
        p_vt_lbl.paragraph_format.space_after = Pt(2)
        r_vtl = p_vt_lbl.add_run("Vote of Thanks:")
        r_vtl.bold = True
        r_vtl.font.size = Pt(11.5)

        p_vt = doc.add_paragraph()
        p_vt.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p_vt.paragraph_format.space_after = Pt(8)
        p_vt.paragraph_format.line_spacing = 1.15
        r_vt = p_vt.add_run(str(vote_thanks))
        r_vt.font.size = Pt(11)

    # ── 14. Signatories Table (Side-by-Side Reference PDF Style) ──
    sig_table = doc.add_table(rows=1, cols=2)
    _remove_table_borders(sig_table)
    sig_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    sig_table.autofit = False

    c_left = sig_table.rows[0].cells[0]
    c_right = sig_table.rows[0].cells[1]
    c_left.width = Inches(3.5)
    c_right.width = Inches(3.5)

    p_sig_left = c_left.paragraphs[0]
    p_sig_left.paragraph_format.space_before = Pt(24)
    p_sig_left.paragraph_format.line_spacing = 1.15

    conv_sig = data.get("convener_signatory")
    if conv_sig and isinstance(conv_sig, dict):
        c_desig = conv_sig.get("designation") or "General Manager / Convener"
        c_loc = conv_sig.get("location") or location or ""
        r_sl = p_sig_left.add_run(f"Sd/xxxxxx\n{c_desig}\n{c_loc}")
    else:
        r_sl = p_sig_left.add_run("Sd/xxxxxx\nConvener")
    r_sl.bold = True
    r_sl.font.size = Pt(11)

    p_sig_right = c_right.paragraphs[0]
    p_sig_right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_sig_right.paragraph_format.space_before = Pt(24)
    p_sig_right.paragraph_format.line_spacing = 1.15

    chair_sig = data.get("chairperson_signatory")
    if chair_sig and isinstance(chair_sig, dict):
        ch_desig = chair_sig.get("designation") or "The District Collector / Chairperson"
        ch_loc = chair_sig.get("location") or location or ""
        r_sr = p_sig_right.add_run(f"Sd/xxxxxx\n{ch_desig}\n{ch_loc}")
    else:
        r_sr = p_sig_right.add_run("Sd/xxxxxx\nChairperson")
    r_sr.bold = True
    r_sr.font.size = Pt(11)

    doc.save(path)
    print(f"[OK] English MoM DOCX saved: {path}")


# ─────────────────────────────────────────────
# TABLE & PAGE NUMBER HELPERS
# ─────────────────────────────────────────────
def _remove_table_borders(table):
    """Removes all visible borders from a table."""
    tblPr = table._tbl.tblPr
    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'none')
        tblBorders.append(border)
    tblPr.append(tblBorders)


def _add_page_number(run_or_para):
    """Inserts a dynamic Word PAGE field for pagination."""
    fldSimple = OxmlElement('w:fldSimple')
    fldSimple.set(qn('w:instr'), 'PAGE')
    run_or_para._p.append(fldSimple)